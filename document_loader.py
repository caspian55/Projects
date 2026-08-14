import os
from typing import List
from pypdf import PdfReader
import docx  
# File loading

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file, page by page."""
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(page_text)
    return "\n\n".join(parts)
def load_docx(file_path: str) -> str:
    """Extract text from a Word (.docx) file, paragraph by paragraph."""
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables in the document
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)
def load_txt(file_path: str) -> str:
    """Read a plain text or markdown file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
def load_document(file_path: str) -> str:
    """Dispatch to the correct loader based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = load_pdf(file_path)
    elif ext == ".docx":
        text = load_docx(file_path)
    elif ext in (".txt", ".md"):
        text = load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not text.strip():
        raise ValueError("No extractable text found in this file (it may be a scanned/image-only document).")

    return text
# Chunking (a small, dependency-free recursive character splitter,
# similar in spirit to LangChain's RecursiveCharacterTextSplitter)
def _merge_splits(splits: List[str], separator: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    chunks = []
    current: List[str] = []
    current_len = 0

    for s in splits:
        s_len = len(s)
        added_len = s_len + (len(separator) if current else 0)

        if current and current_len + added_len > chunk_size:
            chunk = separator.join(current).strip()
            if chunk:
                chunks.append(chunk)

            # Slide the window forward, keeping enough trailing pieces to satisfy the overlap
            while current and current_len > chunk_overlap:
                removed = current.pop(0)
                current_len -= len(removed) + (len(separator) if current else 0)

        current.append(s)
        current_len += s_len + (len(separator) if len(current) > 1 else 0)

    if current:
        chunk = separator.join(current).strip()
        if chunk:
            chunks.append(chunk)

    return chunks
def _recursive_split(text: str, separators: List[str], chunk_size: int, chunk_overlap: int) -> List[str]:
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    if not separators:
        # Last resort: hard cut by character count
        step = max(chunk_size - chunk_overlap, 1)
        return [text[i:i + chunk_size] for i in range(0, len(text), step)]

    sep, rest = separators[0], separators[1:]
    splits = [s for s in text.split(sep)] if sep else list(text)
    splits = [s for s in splits if s != ""]

    good_splits: List[str] = []
    final_chunks: List[str] = []

    for s in splits:
        if len(s) < chunk_size:
            good_splits.append(s)
        else:
            if good_splits:
                final_chunks.extend(_merge_splits(good_splits, sep, chunk_size, chunk_overlap))
                good_splits = []
            final_chunks.extend(_recursive_split(s, rest, chunk_size, chunk_overlap))

    if good_splits:
        final_chunks.extend(_merge_splits(good_splits, sep, chunk_size, chunk_overlap))

    return final_chunks
def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks, preferring to break on paragraph,
    then line, then sentence, then word boundaries.
    """
    text = text.strip()
    if not text:
        return []

    chunk_overlap = min(chunk_overlap, max(chunk_size - 1, 0))
    separators = ["\n\n", "\n", ". ", " "]
    chunks = _recursive_split(text, separators, chunk_size, chunk_overlap)
    return [c.strip() for c in chunks if c.strip()]